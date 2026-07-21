"""Build Korean Word translation of Andy's Summer English Camp proposal."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "Summer_English_Camp_제안서_한글번역.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_ko(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True)
    return p


def add_para(doc, text, size=11, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            run = cells[c_i].paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("여름방학 영어캠프")
    set_run_font(r, size=22, bold=True, color=RGBColor(0x06, 0x4E, 0x3B))

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("초등 프로그램 제안서 (한글 번역본)")
    set_run_font(r, size=14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "원문: Summer Vacation English Camp — Elementary Program Proposal\n"
        "작성: Andrew Regan  ·  제출처: 시청(지자체)  ·  일자: 2026년 7월\n"
        "번역·정리: Tourmaker"
    )
    set_run_font(r, size=10, color=RGBColor(0x64, 0x74, 0x8B))
    doc.add_paragraph()

    note = doc.add_paragraph()
    r = note.add_run(
        "※ 본 문서는 영문 원문의 한글 번역입니다. 원문의 예산·인력 배정란은 공란이었으며, "
        "본 번역본에도 금액·단가는 포함하지 않았습니다."
    )
    set_run_font(r, size=9, color=RGBColor(0x94, 0xA3, 0xB8))

    # 1
    add_heading_ko(doc, "1. 프로그램 개요", 1)
    add_para(
        doc,
        "초등학생 대상 6일(총 60시간) 여름 영어캠프입니다. 오전에는 원어민 교사가 교실에서 "
        "테마 수업을 진행하고, 오후에는 체육관에서 4개 고정 팀의 미니올림픽 팀 대항전을 진행합니다. "
        "일부 오후 이벤트는 오전 수업과 연계되며, 일부는 순수 체육·재미 활동입니다. "
        "누적 점수판으로 캠프 전 기간 순위를 관리하고, 마지막 날은 챔피언십 결승과 시상식으로 마무리합니다.",
    )
    add_para(doc, "핵심 지표", bold=True)
    add_bullets(
        doc,
        [
            "캠프 일수: 6일",
            "고정 팀: 4팀",
            "팀 이벤트: 24개",
            "총 운영시간: 60시간",
        ],
    )

    add_heading_ko(doc, "1.1 팀 및 점수 구조", 2)
    add_bullets(
        doc,
        [
            "1일차에 학생을 4개 고정 팀으로 편성하며, 캠프 종료까지 팀을 유지합니다.",
            "각 팀은 팀 이름을 정하고 깃발·마스코트를 디자인하는 것을 첫 팀빌딩 이벤트로 진행합니다.",
            "매일 오후 3시간 동안 언어 연계 미션과 순수 체육 대항전을 섞어 4개 이벤트를 진행하며, 각 이벤트마다 점수를 부여합니다.",
            "매일 종료 시 누적 점수판을 게시·갱신하여 캠프 전체 기간에 걸쳐 순위 모멘텀을 유지합니다.",
            "최종일에는 미니올림픽 챔피언십과 전체 시상식을 포함합니다.",
        ],
    )

    add_heading_ko(doc, "1.2 일일 시간표 (모든 캠프일에 동일 적용)", 2)
    add_para(
        doc,
        "아래 일정 옵션(평일집중형·주말분산형)과 무관하게 매일 동일한 블록 구조를 적용합니다.",
    )
    add_table(
        doc,
        ["시간", "활동", "세부 내용"],
        [
            ["09:00–10:30", "영어 수업 1", "원어민 교사 · 교실. 일일 테마 핵심 어휘·문법"],
            ["10:30–10:45", "휴식 / 간식", "—"],
            ["10:45–12:00", "영어 수업 2", "원어민 교사 · 교실. 말하기·듣기 연습"],
            ["12:00–13:00", "점심", "—"],
            ["13:00–13:10", "팀 허들 & 테마 공개", "체육관. 팀 집합, 오후 4개 이벤트 예고"],
            ["13:10–13:40", "오후 이벤트 1", "오전 수업 연계 언어 미션. 점수 부여"],
            ["13:40–13:50", "워터 브레이크", "—"],
            ["13:50–14:20", "오후 이벤트 2", "팀 계주 또는 장애물 코스. 점수 부여"],
            ["14:20–14:30", "워터 브레이크", "—"],
            ["14:30–15:00", "오후 이벤트 3", "신체 팀 챌린지(줄다리기, 태그 등). 점수 부여"],
            ["15:00–15:10", "워터 브레이크", "—"],
            ["15:10–15:40", "오후 이벤트 4", "당일 최종 팀 배틀/피날레. 점수 부여"],
            ["15:40–16:00", "점수판 시상", "일일 우승 발표, 누적 점수 갱신·게시"],
        ],
    )

    # 2
    add_heading_ko(doc, "2. 일정 옵션 (택 1)", 1)
    add_para(
        doc,
        "시청은 시설 가용성과 가정 일정에 맞는 캘린더 형식을 선택합니다. "
        "두 옵션 모두 동일하게 60시간 수업·경쟁을 포함합니다.",
    )

    add_heading_ko(doc, "2.1 평일 집중형 — 1주간 월~토 연속", 2)
    add_para(doc, "한 주 동안 6일 연속. 방학 등 한 주에 집중하기를 선호하는 가정에 적합합니다.")
    add_table(
        doc,
        ["요일", "기간", "시간", "테마"],
        [
            ["월요일", "1주차", "09:00–16:00", "캠프에 오신 것을 환영합니다!"],
            ["화요일", "1주차", "09:00–16:00", "동식물과 자연"],
            ["수요일", "1주차", "09:00–16:00", "슈퍼히어로"],
            ["목요일", "1주차", "09:00–16:00", "해적과 보물찾기"],
            ["금요일", "1주차", "09:00–16:00", "여행과 탐험"],
            ["토요일", "1주차", "09:00–16:00", "챔피언십 결승 & 시상의 날"],
        ],
    )

    add_heading_ko(doc, "2.2 주말 분산형 — 3주간 토~일", 2)
    add_para(doc, "3주 주말에 걸쳐 6일 분산. 평일 시설이 어렵거나 분산 일정을 선호하는 가정에 적합합니다.")
    add_table(
        doc,
        ["요일", "기간", "시간", "테마"],
        [
            ["토요일", "주말 1", "09:00–16:00", "캠프에 오신 것을 환영합니다!"],
            ["일요일", "주말 1", "09:00–16:00", "동식물과 자연"],
            ["토요일", "주말 2", "09:00–16:00", "슈퍼히어로"],
            ["일요일", "주말 2", "09:00–16:00", "해적과 보물찾기"],
            ["토요일", "주말 3", "09:00–16:00", "여행과 탐험"],
            ["일요일", "주말 3", "09:00–16:00", "챔피언십 결승 & 시상의 날"],
        ],
    )

    # 3
    add_heading_ko(doc, "3. 일일 테마 및 오후 이벤트 라인업", 1)
    add_para(
        doc,
        "각 날은 아동 친화적 테마와 오후 4개 이벤트로 구성됩니다. "
        "3시간 블록을 단일 반복 게임이 아닌 다양한 경쟁으로 채웁니다.",
    )
    themes = [
        ("Day 1", "캠프에 오신 것을 환영합니다!", "인사·자기소개", [
            "팀명 & 깃발 디자인 콘테스트(창의성 점수)",
            "아이스브레이커 계주",
            "휴먼 노트 신뢰 챌린지",
            "입문 깃발뺏기(Capture-the-Flag)",
        ]),
        ("Day 2", "동식물과 자연", "동물 어휘·소리", [
            "동물 제스처 계주(Charades)",
            "정글 크롤 장애물 코스",
            "동물 소리 보물찾기",
            "포식자·피식자 태그 토너먼트",
        ]),
        ("Day 3", "슈퍼히어로", "동작 동사·능력(can/can't)", [
            "히어로 트레이닝 부트캠프(장애물)",
            "‘도시 구하기’ 구조 계주",
            "빌런 피구 토너먼트",
            "슈퍼파워 퀴즈 쇼다운(영어 복습)",
        ]),
        ("Day 4", "해적과 보물찾기", "방향·전치사(지도/단서 읽기)", [
            "보물지도 스캐빈저 헌트",
            "널판 걷기(Walk-the-Plank) 균형 계주",
            "해적선 줄다리기",
            "보물 쟁탈 팀 배틀",
        ]),
        ("Day 5", "여행과 탐험", "장소·교통수단", [
            "여권 퀴즈 챌린지(영어 복습)",
            "세계일주 스테이션 계주",
            "나침반·지도 오리엔티어링",
            "팀 스캐빈저 헌트 피날레",
        ]),
        ("Day 6", "챔피언십 결승 & 시상의 날", "주간 어휘 누적 복습", [
            "클로징 미니올림픽 챔피언십(주간 베스트 이벤트 재경기)",
            "팀 대 팀 챔피언십 계주",
            "최종 점수판 공개",
            "시상식·수료증·팀 사진",
        ]),
    ]
    for day, theme, focus, events in themes:
        add_para(doc, f"{day} — {theme}", bold=True, size=12)
        add_para(doc, f"언어 초점: {focus}", size=10)
        add_bullets(doc, events)

    # 4
    add_heading_ko(doc, "4. 일별 필요 자재", 1)
    add_table(
        doc,
        ["일차", "오전 수업 자재", "오후 이벤트 자재"],
        [
            ["1", "이름표, 환영·인사 플래시카드, 출석표", "포스터지·마카, 콘·계주봉, 조끼/깃발"],
            ["2", "동물 플래시카드, 동물 소리 오디오, 워크시트", "제스처 카드, 콘·매트·후프, 숨김용 그림 카드, 태그 밴드"],
            ["3", "히어로 어휘 카드, 만화형 워크시트", "콘·매트·터널, 빈백·구조 매트, 폼 피구공, 퀴즈 보드"],
            ["4", "해적 어휘 카드, 보물지도 워크시트", "인쇄 지도·단서, 균형선 테이프·빈백, 줄다리기 로프, 보물 토큰"],
            ["5", "교통수단 플래시카드, 여권 워크시트", "여권 북렛·스탬프, 스테이션 소품, 나침반·지도 카드, 단서 카드"],
            ["6", "누적 복습 워크시트, 어휘 복습 게임", "이전 일자 장비 재사용, 점수판, 수료증·메달/트로피"],
        ],
    )

    # 5
    add_heading_ko(doc, "5. 오후 활동 설명 요약", 1)

    activities = {
        "Day 1 — 캠프에 오신 것을 환영합니다!": [
            ("팀명 & 깃발 디자인 콘테스트", "팀이 이름을 정하고 포스터 크기 깃발을 꾸밉니다. 속도보다 창의성·협동을 평가합니다."),
            ("아이스브레이커 계주", "콘까지 왕복 후 다음 동료를 태그. 가장 빠른 팀이 1위 점수."),
            ("휴먼 노트 신뢰 챌린지", "원이 되어 서로 다른 두 동료와 손을 잡은 뒤, 손을 놓지 않고 한 원으로 풀어냅니다."),
            ("입문 깃발뺏기", "체육관을 반으로 나누고 상대 깃발을 가져오되, 적진에서 태그당하면 실패."),
        ],
        "Day 2 — 동식물과 자연": [
            ("동물 제스처 계주", "한 명씩 카드 동물을 몸짓으로 표현, 팀이 영어로 맞히면 다음 주자."),
            ("정글 크롤 장애물", "매트·후프·콘으로 만든 정글 코스를 기어·넘으며 완주 경쟁."),
            ("동물 소리 보물찾기", "숨겨진 동물 카드를 찾고 해당 소리를 내어 점수 획득."),
            ("포식자·피식자 태그", "포식자가 제한 시간 내 피식자를 태그. 라운드마다 역할 교대."),
        ],
        "Day 3 — 슈퍼히어로": [
            ("히어로 트레이닝 부트캠프", "기어가기·점프·균형 등 타임드 장애물. 팀 합산 시간으로 점수."),
            ("도시 구하기 구조 계주", "빈백 ‘시민’을 안전지대로 옮기고 돌아와 다음 구조자를 태그."),
            ("빌런 피구 토너먼트", "폼공 소프트 피구. 마지막까지 남은 팀이 라운드 승리."),
            ("슈퍼파워 퀴즈 쇼다운", "당일 수업 영어 복습 문제. 정답 시 점수."),
        ],
        "Day 4 — 해적과 보물찾기": [
            ("보물지도 스캐빈저 헌트", "영어 방향 단서가 있는 지도로 숨겨진 ‘보물’ 토큰을 찾습니다."),
            ("널판 걷기 균형 계주", "테이핑 라인 위를 빈백을 떨어뜨리지 않고 걸어 다음 주자에게 전달."),
            ("해적선 줄다리기", "로프 양끝을 당겨 상대를 중앙선 너머로 끌어당기면 승리."),
            ("보물 쟁탈 팀 배틀", "깃발뺏기와 유사하나, 상대 홈베이스의 보물 토큰을 탈취."),
        ],
        "Day 5 — 여행과 탐험": [
            ("여권 퀴즈 챌린지", "장소·여행 어휘 복습. 정답마다 ‘여권 스탬프’."),
            ("세계일주 스테이션 계주", "나라 테마 스테이션을 돌며 짧은 신체 과제를 수행."),
            ("나침반·지도 오리엔티어링", "간단한 나침반·체육관 지도로 체크포인트를 순서대로 찾고 방향 어휘 연습."),
            ("팀 스캐빈저 헌트 피날레", "주간 테마 단서를 모아 마지막 숨김 아이템을 찾습니다."),
        ],
        "Day 6 — 챔피언십 결승 & 시상의 날": [
            ("클로징 미니올림픽 챔피언십", "스태프 선정 주간 베스트 이벤트를 상위 2팀이 결승으로 재경기."),
            ("팀 대 팀 챔피언십 계주", "상위 2팀 헤드투헤드 계주. 나머지 학생은 응원."),
            ("최종 점수판 공개", "전 이벤트 누적 점수를 집계해 최종 순위 발표."),
            ("시상식·수료증·팀 사진", "전원 수료증, 상위팀 트로피/메달, 단체 사진."),
        ],
    }
    for title, pairs in activities.items():
        add_para(doc, title, bold=True, size=12)
        for name, desc in pairs:
            add_para(doc, f"• {name}: {desc}", size=10, space_after=3)

    # 6
    add_heading_ko(doc, "6. 인력 편성 (원문 담당자란은 공란)", 1)
    add_para(doc, "원문 제안서의 Assigned To(담당자)란은 비어 있었습니다. 역할과 책임만 번역·정리합니다.")
    add_table(
        doc,
        ["역할", "주요 책임"],
        [
            ["원어민 영어교사", "오전 교실 수업, 오후 영어 연계 미션 지원"],
            ["캠프 코디네이터/총괄", "일정·안전·점수판·학부모 커뮤니케이션 총괄"],
            ["팀 리더/활동 보조", "팀별 인솔, 오후 이벤트 운영 보조, 안전 관찰"],
        ],
    )

    # 7
    add_heading_ko(doc, "7. 예산 항목 (원문 금액 공란)", 1)
    add_para(
        doc,
        "원문에는 아래 카테고리만 제시되고 Estimated Cost는 비어 있었습니다. "
        "본 번역본에는 금액을 기입하지 않습니다. 세부 산출은 별도 협의 문서에서 다룹니다.",
    )
    add_bullets(
        doc,
        [
            "원어민 교사",
            "시설 대관(교실·체육관)",
            "교재·소모품",
            "팀 이벤트 장비",
            "시상·수료증·상품",
            "잡비/예비비",
            "합계",
        ],
    )

    add_heading_ko(doc, "맺음말", 1)
    add_para(doc, "본 제안서 검토에 시간을 내어 주셔서 감사합니다.")
    add_para(
        doc,
        "— End of translation —\nSource: Summer_English_Camp_Schedule_Proposal.docx (Andrew Regan, July 2026)",
        size=9,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
